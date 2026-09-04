"""
extractor.py
多格式文档 → 统一文本流。
- PDF     ：PyMuPDF (fitz)，按页提取，保留 page_number
- Word    ：python-docx，段落 + 表格
- Excel   ：pandas，逐工作表转为文本表格
- Markdown/纯文本：直接读取
所有格式输出统一的 ExtractedPage 列表。
"""

from dataclasses import dataclass
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".md", ".txt"}


@dataclass
class ExtractedPage:
    page_number: int  # 1-based；无分页概念的非 PDF 文档为 0
    text: str


def _extract_pdf(path: Path):
    import fitz  # PyMuPDF

    pages = []
    with fitz.open(str(path)) as doc:
        for i, page in enumerate(doc, start=1):
            text = page.get_text().strip()
            if text:
                pages.append(ExtractedPage(page_number=i, text=text))
    return pages


def _extract_docx(path: Path):
    from docx import Document

    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return [ExtractedPage(page_number=0, text=text)] if text.strip() else []


def _extract_xlsx(path: Path):
    import pandas as pd

    sheets = pd.read_excel(str(path), sheet_name=None)
    blocks = []
    for sheet_name, df in sheets.items():
        lines = [f"【工作表：{sheet_name}】"]
        cols = [str(c) for c in df.columns]
        lines.append(" | ".join(cols))
        for _, row in df.iterrows():
            parts = []
            for c in df.columns:
                v = row[c]
                parts.append(f"{c}: {'' if pd.isna(v) else v}")
            lines.append(" | ".join(parts))
        blocks.append("\n".join(lines))
    text = "\n\n".join(blocks)
    return [ExtractedPage(page_number=0, text=text)] if text.strip() else []


def _extract_markdown(path: Path):
    text = path.read_text(encoding="utf-8").strip()
    return [ExtractedPage(page_number=0, text=text)] if text else []


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
    ".md": _extract_markdown,
    ".txt": _extract_markdown,
}


def extract_document(path) -> list:
    """提取文档为统一文本流（ExtractedPage 列表）。空文档抛 ValueError。"""
    p = Path(path)
    ext = p.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件格式: {ext}（支持 {', '.join(sorted(SUPPORTED_EXTENSIONS))}）"
        )
    pages = _EXTRACTORS[ext](p)
    pages = [pg for pg in pages if pg.text.strip()]
    if not pages:
        raise ValueError(f"未能从文档 {p.name} 中提取到任何文本内容")
    return pages
